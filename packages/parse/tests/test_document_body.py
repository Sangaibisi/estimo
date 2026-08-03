"""S12-2: the persisted document body — its budget, and the join it must not break.

Both properties here were WRONG in the first cut and are exactly what the adversarial
review reproduced, so each test states the failure it prevents.
"""

from pathlib import Path

from docx import Document as create_document
from estimo_parse import parse_brd
from estimo_parse.models import MAX_BLOCK_CHARS, MAX_BODY_BYTES


def _write(tmp_path: Path, name: str, build) -> Path:  # type: ignore[no-untyped-def]
    doc = create_document()
    build(doc)
    path = tmp_path / name
    doc.save(str(path))
    return path


def _body_bytes(parsed) -> int:  # type: ignore[no-untyped-def]
    return sum(len(block.model_dump_json().encode()) for block in parsed.blocks)


def test_the_budget_bounds_what_is_actually_stored(tmp_path: Path) -> None:
    """The first cut charged block TEXT, so the JSON it stored ran 3.6x over on this
    repo's own fixture and ~545x on many short paragraphs under a deep heading trail:
    a 12 KB .docx could persist an 18 MB row while the budget reported itself met.
    """

    def build(doc) -> None:  # type: ignore[no-untyped-def]
        doc.add_heading("Uzun Bölüm Başlığı — Kapsam ve Kısıtlar Bölümü", level=1)
        doc.add_heading("Alt Başlık — Detaylı İnceleme ve Değerlendirme", level=2)
        for index in range(6000):
            doc.add_paragraph(f"S{index}")

    parsed = parse_brd(_write(tmp_path, "many-tiny-paragraphs.docx", build))
    assert parsed.body_truncated is True
    stored = _body_bytes(parsed)
    assert stored <= MAX_BODY_BYTES * 1.05, (
        f"stored {stored} bytes against a {MAX_BODY_BYTES} budget"
    )


def test_a_requirement_always_has_the_block_it_points_at(tmp_path: Path) -> None:
    """A BRD puts its requirements at the END, so cutting the body at a prefix
    orphaned precisely the rows the Reading Room exists for: clicking one scrolled
    nowhere and highlighted nothing, with no message and no disabled state.
    """

    def build(doc) -> None:  # type: ignore[no-untyped-def]
        doc.add_heading("Giriş", level=1)
        for index in range(4000):
            doc.add_paragraph(f"Bağlam paragrafı {index} — kapsam dışı açıklama metni.")
        doc.add_heading("3. İş Gereksinimleri", level=1)
        doc.add_paragraph("G-01: Abone taksitli ödeme seçeneğini görebilmelidir.")
        doc.add_paragraph("G-02: Taksit tutarı faturaya ayrı kalem olarak yansımalıdır.")

    parsed = parse_brd(_write(tmp_path, "requirements-at-the-end.docx", build))
    assert parsed.body_truncated is True, "this document must exceed the budget"
    assert parsed.requirements, "the fixture produced no requirements"

    kept = {block.source_ref for block in parsed.blocks}
    orphans = [req.id for req in parsed.requirements if req.source_ref not in kept]
    assert not orphans, f"requirements point at blocks that were dropped: {orphans}"


def test_one_enormous_block_does_not_end_the_document(tmp_path: Path) -> None:
    """Returning at the first over-budget block meant a single huge table could yield
    an EMPTY body — and the pane then told the reader to re-upload a document it had
    parsed perfectly well."""

    def build(doc) -> None:  # type: ignore[no-untyped-def]
        doc.add_paragraph("x" * (MAX_BLOCK_CHARS * 4))
        doc.add_heading("3. İş Gereksinimleri", level=1)
        doc.add_paragraph("G-01: Abone taksitli ödeme seçeneğini görebilmelidir.")

    parsed = parse_brd(_write(tmp_path, "one-huge-block.docx", build))
    assert parsed.blocks, "an oversized first block emptied the whole body"
    giant = next(block for block in parsed.blocks if block.index == 0)
    assert len(giant.text) == MAX_BLOCK_CHARS and giant.text_truncated is True
    assert any(block.kind == "heading" for block in parsed.blocks), "the tail was lost"


def test_a_normal_brd_is_kept_whole(tmp_path: Path) -> None:
    """The guard must not cost anything on documents of ordinary size."""

    def build(doc) -> None:  # type: ignore[no-untyped-def]
        doc.add_heading("3. İş Gereksinimleri", level=1)
        for index in range(1, 9):
            doc.add_paragraph(f"G-0{index}: Gereksinim metni {index}.")

    parsed = parse_brd(_write(tmp_path, "ordinary.docx", build))
    assert parsed.body_truncated is False
    assert not any(block.text_truncated for block in parsed.blocks)
