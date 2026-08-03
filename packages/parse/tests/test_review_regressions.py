"""Regressions for the confirmed findings of the S2 adversarial review.

Each test rebuilds the reviewer's failure scenario as a synthetic docx and asserts the
fixed behavior — if a fix regresses, the original finding fires again here.
"""

from collections.abc import Callable
from pathlib import Path

import pytest
from docx import Document
from estimo_parse import parse_brd, rule_score
from estimo_parse.anchors import detect_anchors

from estimo_core import Requirement


def build_docx(tmp_path: Path, name: str, build: Callable[[Document], None]) -> Path:
    doc = Document()
    build(doc)
    path = tmp_path / name
    doc.save(path)
    return path


def add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    for j, cell in enumerate(header):
        table.rows[0].cells[j].text = cell
    for i, row in enumerate(rows, start=1):
        for j, cell in enumerate(row):
            table.rows[i].cells[j].text = cell


def test_meta_table_anchors_are_quarantined(tmp_path: Path) -> None:
    """Finding 1 (critical): anchors in table cells must not smuggle past quarantine."""

    def build(doc: Document) -> None:
        doc.add_heading("Talep Özeti", level=0)
        add_table(
            doc,
            ["Alan", "Değer"],
            [
                ["Bütçe", "Bu iş için ayrılan bütçe azami 40 adam-gündür."],
                ["Hedef", "15 Ekim 2026 tarihinde canlıya alınması planlanmaktadır."],
                ["Not", "Kapsam işletme ile netleştirilecektir."],
            ],
        )
        doc.add_heading("Gereksinimler", level=1)
        doc.add_paragraph("G-01: Fatura ekranında bakiye gösterilmelidir.", style="List Bullet")

    parsed = parse_brd(build_docx(tmp_path, "meta-anchor.docx", build))
    types = {a.type for a in parsed.doc_anchors}
    assert "budget" in types
    assert "deadline" in types
    assert any("netleştirilecek" in p for p in parsed.open_points)


def test_revision_table_is_not_a_requirement_table(tmp_path: Path) -> None:
    """Finding 2 (critical): 'Gereksinim Sahibi' columns / revision tables must not
    mint requirements nor starve the heuristic rung."""

    def build(doc: Document) -> None:
        doc.add_heading("Bayi Bildirim Talebi", level=0)
        doc.add_heading("Revizyon Geçmişi", level=1)
        add_table(
            doc,
            ["Versiyon", "Tarih", "Gereksinim Sahibi"],
            [["1.0", "01.07.2026", "K. Arslan"], ["1.1", "10.07.2026", "E. Şahin"]],
        )
        doc.add_heading("Talep", level=1)
        doc.add_paragraph("Bayilere sipariş durumu bildirimi gönderilmelidir.")

    parsed = parse_brd(build_docx(tmp_path, "revision.docx", build))
    assert all(r.extraction == "heuristic" for r in parsed.requirements)
    assert len(parsed.requirements) == 1
    assert not any("Arslan" in r.text for r in parsed.requirements)


def test_annex_label_does_not_starve_heuristics(tmp_path: Path) -> None:
    """Finding 3 (major): 'EK-1:' style labels are not coded requirements."""

    def build(doc: Document) -> None:
        doc.add_heading("Kısa Talep", level=0)
        doc.add_paragraph("EK-1: Ekran görüntüleri ektedir.")
        doc.add_paragraph("Abone geçmişi ekranında son 12 ay listelenmelidir.")
        doc.add_paragraph("Rapor çıktısı Excel olarak indirilebilir olmalı.")

    parsed = parse_brd(build_docx(tmp_path, "annex.docx", build))
    assert len(parsed.requirements) == 2
    assert all(r.id.startswith("REQ-H") for r in parsed.requirements)
    assert not any(r.id == "REQ-EK-1" for r in parsed.requirements)


def test_coded_and_table_restatements_deduplicate(tmp_path: Path) -> None:
    """Finding 4 (major): the same code in a paragraph and a requirement table must
    collapse to one record, preferring the acceptance-carrying table copy."""

    def build(doc: Document) -> None:
        doc.add_heading("Talep", level=0)
        doc.add_paragraph("G-01: Konsolide fatura talebi alınabilmelidir.", style="List Bullet")
        doc.add_heading("Detay Tablosu", level=1)
        add_table(
            doc,
            ["No", "Gereksinim", "Kabul Kriteri"],
            [["G-01", "Konsolide fatura talebi alınabilmelidir.", "İlk dönemde tek fatura."]],
        )

    parsed = parse_brd(build_docx(tmp_path, "dedup.docx", build))
    assert [r.id for r in parsed.requirements] == ["REQ-G-01"]
    only = parsed.requirements[0]
    assert only.extraction == "table"
    assert only.acceptance_criteria == "İlk dönemde tek fatura."


def test_rule_score_is_turkish_case_insensitive() -> None:
    """Finding 5 (major): ALL-CAPS Turkish (İ/I) must score like lowercase."""
    lower = Requirement(
        id="REQ-T-1", text="Kurumsal müşteriler için farklı koşullar uygulanabilir."
    )
    upper = Requirement(
        id="REQ-T-2", text="KURUMSAL MÜŞTERİLER İÇİN FARKLI KOŞULLAR UYGULANABİLİR."
    )
    assert rule_score(upper)[0] == pytest.approx(rule_score(lower)[0])
    assert rule_score(upper)[1] == rule_score(lower)[1]


class TestInvertedAnchorOrders:
    """Finding 6 (major): keyword/number order must not matter within a sentence."""

    def test_number_first_budget(self) -> None:
        anchors = detect_anchors("Çalışma 90 adam-günü aşmayacak şekilde bütçelendi.")
        assert [a.type for a in anchors] == ["budget"]

    def test_keyword_first_deadline(self) -> None:
        anchors = detect_anchors("Canlıya geçişin 15 Ekim 2026 tarihinde yapılması beklenir.")
        assert [a.type for a in anchors] == ["deadline"]

    def test_original_orders_still_detected(self) -> None:
        assert [a.type for a in detect_anchors("Ayrılan bütçe azami 90 adam-gündür.")] == ["budget"]
        assert [
            a.type for a in detect_anchors("Kampanya lansmanı 30 Eylül 2026 tarihine yetişmelidir.")
        ] == ["deadline"]
