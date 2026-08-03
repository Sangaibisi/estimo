"""S6 behavior: calibration, bands, estimator flow, critic, BoE rendering, LOO eval."""

from pathlib import Path

import pytest
import respx
from docx import Document as open_document
from estimo_estimate import (
    band_from_analogs,
    estimate_state,
    leave_one_out,
    render_boe_docx,
    review_boe,
)
from estimo_estimate.bands import SMALL_ITEM_MIN_SPREAD
from estimo_estimate.calibration import ErrorDistribution
from estimo_estimate.estimator import _llm_adjust_likely
from estimo_estimate.prompts import load_estimate_prompt
from estimo_knowledge import AnalogyCard, import_seed
from estimo_pipeline import run_brd
from pydantic import SecretStr
from sqlalchemy.ext.asyncio import AsyncSession

from estimo_core import Confidence, ThreePoint
from estimo_gateway import GatewayClient, GatewayConfig

REPO_ROOT = Path(__file__).resolve().parents[3]
FIXTURES = REPO_ROOT / "fixtures" / "brd"
SEED = REPO_ROOT / "fixtures" / "seed" / "sample-seed.csv"
BASE_URL = "http://gateway.local/v1"

DIST = ErrorDistribution(q10=0.8, q50=1.1, q90=1.5, samples=20, prior_based=False)


def card(entry_id: str, actual: float | None, likely: float | None = None) -> AnalogyCard:
    import uuid

    return AnalogyCard(
        entry_id=uuid.uuid5(uuid.NAMESPACE_URL, entry_id),
        rank=1,
        brd_ref="B",
        item_title=entry_id,
        estimate=ThreePoint(
            optimistic=likely or 4, likely=likely or 5, pessimistic=(likely or 5) * 2
        )
        if likely
        else None,
        actual_effort=actual,
        actual_source="timesheet" if actual is not None else None,
    )


class TestBands:
    def test_prompt_is_versioned(self) -> None:
        assert load_estimate_prompt().id == "estimate-v1"

    def test_no_analogs_returns_none(self) -> None:
        assert band_from_analogs([], DIST) is None

    def test_actuals_drive_likely_and_quantiles_drive_width(self) -> None:
        band = band_from_analogs([card("a", 10.0), card("b", 12.0), card("c", 8.0)], DIST)
        assert band is not None
        # All three points derive from the analog MEDIAN (10.0) — q50 applies once.
        assert band.range.likely == pytest.approx(10.0 * 1.1)
        assert band.range.optimistic == pytest.approx(10.0 * 0.8)
        assert band.range.pessimistic == pytest.approx(10.0 * 1.5)
        assert band.confidence is Confidence.HIGH

    def test_small_item_floor(self) -> None:
        band = band_from_analogs([card("a", 1.0), card("b", 1.0)], DIST)
        assert band is not None
        assert band.range.pessimistic >= band.range.likely + SMALL_ITEM_MIN_SPREAD

    def test_prior_distribution_caps_confidence(self) -> None:
        band = band_from_analogs(
            [card("a", 10.0), card("b", 9.0)], ErrorDistribution.prior(samples=3)
        )
        assert band is not None
        assert band.confidence is Confidence.MEDIUM


class TestLlmAdjust:
    @respx.mock
    async def test_out_of_band_llm_value_rejected_and_anchors_redacted(self) -> None:
        route = respx.post(f"{BASE_URL}/chat/completions").respond(
            200,
            json={
                "id": "c",
                "object": "chat.completion",
                "created": 1,
                "model": "balanced",
                "choices": [
                    {
                        "index": 0,
                        "message": {"role": "assistant", "content": '{"likely": 99.0}'},
                        "finish_reason": "stop",
                    }
                ],
                "usage": {"prompt_tokens": 1, "completion_tokens": 1, "total_tokens": 2},
            },
        )
        client = GatewayClient(
            GatewayConfig(
                base_url=BASE_URL,
                api_key=SecretStr("sk-test"),
                profiles={"default": "balanced"},
            )
        )
        band = band_from_analogs([card("a", 10.0), card("b", 11.0)], DIST)
        assert band is not None
        try:
            likely, _ = await _llm_adjust_likely(
                client,
                load_estimate_prompt(),
                "Taksit kalemi. Bu iş için ayrılan bütçe azami 90 adam-gündür.",
                band,
            )
        finally:
            await client.aclose()
        assert likely == band.range.likely  # 99 is outside the band -> draft kept
        payload = route.calls.last.request.content.decode("utf-8")
        assert "90 adam-gün" not in payload  # PRINCIPLES #5 at the estimate boundary


@pytest.mark.db
class TestEstimatorFlow:
    @pytest.fixture
    async def seeded(self, session: AsyncSession, clean_tables: None) -> AsyncSession:
        report = await import_seed(session, SEED)
        assert report.rejected == []
        return session

    async def test_clean_brd_yields_full_boe(self, seeded: AsyncSession) -> None:
        state = await run_brd(FIXTURES / "BRD-AUR-26-02-konsolide-fatura.docx")
        boe = await estimate_state(seeded, state)
        assert len(boe.lines) == 6
        assert all(line.evidence for line in boe.lines)
        assert any(ref.uri.startswith("ledger://") for line in boe.lines for ref in line.evidence)
        assert boe.total.likely > 0
        assert review_boe(boe, state) == []

    async def test_awaiting_answers_state_refused(self, seeded: AsyncSession) -> None:
        state = await run_brd(FIXTURES / "BRD-AUR-26-01-taksitlendirme.docx")
        with pytest.raises(ValueError, match="ready_for_estimation"):
            await estimate_state(seeded, state)

    async def test_boe_docx_renders_turkish(self, seeded: AsyncSession, tmp_path: Path) -> None:
        state = await run_brd(FIXTURES / "BRD-AUR-26-02-konsolide-fatura.docx")
        boe = await estimate_state(seeded, state)
        out = tmp_path / "boe.docx"
        render_boe_docx(boe, out)
        doc = open_document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Taslak Efor Dokümanı" in text
        assert "TASLAK" in text
        assert len(doc.tables) >= 2  # items + signatures
        likely_cell = doc.tables[0].rows[1].cells[2].text
        assert "," in likely_cell  # TR decimal comma

    async def test_leave_one_out_eval(self, seeded: AsyncSession) -> None:
        result = await leave_one_out(seeded)
        assert result.cases >= 10
        assert result.coverage >= 0.5, f"coverage {result.coverage:.0%}"
        assert result.mae <= result.naive_mae + 0.5, (
            f"mae {result.mae:.2f} vs naive {result.naive_mae:.2f}"
        )


class TestS6ReviewRegressions:
    def test_band_edges_derive_from_median_not_likely(self) -> None:
        """Finding 1 (critical): q50 must not apply twice."""
        dist = ErrorDistribution(q10=1.2, q50=1.5, q90=2.0, samples=20, prior_based=False)
        band = band_from_analogs([card("a", 10.0), card("b", 10.0)], dist)
        assert band is not None
        assert (band.range.optimistic, band.range.likely, band.range.pessimistic) == (
            12.0,
            15.0,
            20.0,
        )

    @pytest.mark.db
    async def test_embed_boundary_sees_redacted_text(
        self, session: AsyncSession, clean_tables: None
    ) -> None:
        """Finding 2 (major): anchors must not reach client.embed."""
        import json as jsonlib

        from estimo_pipeline import PipelineState

        from estimo_core import Requirement, WorkItem

        await import_seed(session, SEED)
        item = WorkItem(
            id="WI-X",
            brd_ref="B-X",
            title="Taksit kalemi eklenmesi",
            description="Bu iş için ayrılan bütçe azami 90 adam-gündür.",
            requirement_ids=("REQ-X",),
            module_tags=("billing-core",),
        )
        state = PipelineState(
            source_path="x",
            requirements=(Requirement(id="REQ-X", text="Taksit kalemi eklenmelidir."),),
            work_items=(item,),
            status="ready_for_estimation",
        )
        with respx.mock:
            embed_route = respx.post(f"{BASE_URL}/embeddings").respond(
                200,
                json={
                    "object": "list",
                    "model": "embed",
                    "data": [{"object": "embedding", "index": 0, "embedding": [0.1] * 4}],
                    "usage": {"prompt_tokens": 1, "total_tokens": 1},
                },
            )
            respx.post(f"{BASE_URL}/chat/completions").respond(
                200,
                json={
                    "id": "c",
                    "object": "chat.completion",
                    "created": 1,
                    "model": "balanced",
                    "choices": [
                        {
                            "index": 0,
                            "message": {"role": "assistant", "content": "{}"},
                            "finish_reason": "stop",
                        }
                    ],
                    "usage": {"prompt_tokens": 1, "completion_tokens": 1, "total_tokens": 2},
                },
            )
            client = GatewayClient(
                GatewayConfig(
                    base_url=BASE_URL,
                    api_key=SecretStr("sk-test"),
                    profiles={"default": "balanced"},
                )
            )
            try:
                await estimate_state(session, state, client=client)
            finally:
                await client.aclose()
            assert embed_route.called
            body = jsonlib.loads(embed_route.calls.last.request.content)
            joined = " ".join(body["input"]) if isinstance(body["input"], list) else body["input"]
            assert "90 adam-gün" not in joined
            assert "budget-karantina" in joined

    @pytest.mark.db
    async def test_no_analog_lines_use_note_scheme(
        self, session: AsyncSession, clean_tables: None
    ) -> None:
        """Finding 3 (major): empty ledger yields note:// annotations, never fake ledger://."""
        from estimo_pipeline import PipelineState

        from estimo_core import Requirement, WorkItem

        state = PipelineState(
            source_path="x",
            requirements=(Requirement(id="REQ-Y", text="Yeni rapor ekranı geliştirilmelidir."),),
            work_items=(
                WorkItem(
                    id="WI-Y",
                    brd_ref="B-Y",
                    title="Yeni rapor ekranı",
                    requirement_ids=("REQ-Y",),
                ),
            ),
            status="ready_for_estimation",
        )
        boe = await estimate_state(session, state)
        uris = [ref.uri for line in boe.lines for ref in line.evidence]
        assert not any(uri.startswith("ledger://") for uri in uris)
        assert any(uri.startswith("note://") for uri in uris)
        assert boe.lines[0].confidence is Confidence.LOW

    @pytest.mark.db
    async def test_answered_questions_not_listed_open_in_docx(
        self, session: AsyncSession, clean_tables: None, tmp_path: Path
    ) -> None:
        """Finding 4 (minor): applied answers leave the open-questions section."""
        from estimo_pipeline import resume_with_answers

        await import_seed(session, SEED)
        state = await run_brd(FIXTURES / "BRD-AUR-26-01-taksitlendirme.docx")
        resumed = await resume_with_answers(
            state,
            {
                "Q-REQ-G-04": (
                    "Kurumsal müşterilerde taksit sayısı 6 ile sınırlıdır ve komisyon "
                    "oranı %5 sabittir."
                )
            },
        )
        boe = await estimate_state(session, resumed)
        out = tmp_path / "boe.docx"
        render_boe_docx(boe, out)
        doc = open_document(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Açık Sorular" not in text or "Q-REQ-G-04" not in text.split("Açık Sorular")[-1]
