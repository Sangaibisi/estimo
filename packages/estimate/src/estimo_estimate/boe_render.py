"""BoE `.docx` rendering (S6-6): the commercial artifact, locale-aware.

The BoE for a Turkish operator renders in Turkish with TR number formatting; the
document carries the full professional anatomy (RESEARCH §4): scope basis, itemized
three-point table, assumption and risk registers, open questions, cone stage,
provenance appendix and signature blocks.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as create_document
from docx.document import Document

from estimo_core import BoeDocument

_LABELS = {
    "tr": {
        "title": "Taslak Efor Dokümanı (Basis of Estimate)",
        "brd": "İlgili BRD",
        "cone": "Belirsizlik Konisi Aşaması",
        "cone_concept": "Konsept (±4x)",
        "cone_approved": "Onaylı Kapsam (±1.5x)",
        "cone_detailed": "Detaylı (±1.25x)",
        "items": "Kalem Tablosu",
        "item": "Kalem",
        "optimistic": "İyimser",
        "likely": "Olası",
        "pessimistic": "Kötümser",
        "confidence": "Güven",
        "evidence": "Kanıt",
        "total": "Toplam",
        "assumptions": "Varsayım Sicili",
        "risks": "Risk Sicili",
        "contingency": "Kontenjan",
        "questions": "Açık Sorular",
        "provenance": "Kanıt Eki (Provenance)",
        "signatures": "İmzalar",
        "name": "Ad Soyad",
        "role": "Rol",
        "scope": "Kapsam",
        "date": "Tarih / İmza",
        "unit": "a-g",
        "draft_note": (
            "Bu doküman Estimo tarafından üretilmiş bir TASLAKTIR; her satır insan "
            "onayı gerektirir."
        ),
    },
    "en": {
        "title": "Basis of Estimate (Draft)",
        "brd": "Source BRD",
        "cone": "Cone of Uncertainty Stage",
        "cone_concept": "Concept (±4x)",
        "cone_approved": "Approved Scope (±1.5x)",
        "cone_detailed": "Detailed (±1.25x)",
        "items": "Estimate Lines",
        "item": "Item",
        "optimistic": "Optimistic",
        "likely": "Likely",
        "pessimistic": "Pessimistic",
        "confidence": "Confidence",
        "evidence": "Evidence",
        "total": "Total",
        "assumptions": "Assumption Register",
        "risks": "Risk Register",
        "contingency": "Contingency",
        "questions": "Open Questions",
        "provenance": "Provenance Appendix",
        "signatures": "Signatures",
        "name": "Name",
        "role": "Role",
        "scope": "Scope",
        "date": "Date / Signature",
        "unit": "pd",
        "draft_note": "This is an Estimo-generated DRAFT; every line requires human sign-off.",
    },
}

_CONE_KEY = {
    "concept": "cone_concept",
    "approved_scope": "cone_approved",
    "detailed": "cone_detailed",
}


def _fmt(value: float, locale: str) -> str:
    text = f"{value:,.1f}"
    if locale == "tr":
        text = text.replace(",", "§").replace(".", ",").replace("§", ".")
    return text


def render_boe_docx(boe: BoeDocument, path: Path) -> None:
    labels = _LABELS["tr"] if boe.locale == "tr" else _LABELS["en"]
    locale = boe.locale
    doc: Document = create_document()

    doc.add_heading(labels["title"], level=0)
    doc.add_paragraph(f"{labels['brd']}: {boe.brd_ref} — {boe.title}")
    doc.add_paragraph(f"{labels['cone']}: {labels[_CONE_KEY[boe.cone_stage.value]]}")
    doc.add_paragraph(labels["draft_note"])

    doc.add_heading(labels["items"], level=1)
    table = doc.add_table(rows=1 + len(boe.lines) + 1, cols=6)
    table.style = "Table Grid"
    header = (
        labels["item"],
        labels["optimistic"],
        labels["likely"],
        labels["pessimistic"],
        labels["confidence"],
        labels["evidence"],
    )
    for column, text in enumerate(header):
        table.rows[0].cells[column].text = text
    for row, line in enumerate(boe.lines, start=1):
        cells = table.rows[row].cells
        cells[0].text = line.work_item_id
        cells[1].text = _fmt(line.range.optimistic, locale)
        cells[2].text = _fmt(line.range.likely, locale)
        cells[3].text = _fmt(line.range.pessimistic, locale)
        cells[4].text = line.confidence.value
        cells[5].text = str(sum(1 for ref in line.evidence if ref.kind.value != "note"))
    total_cells = table.rows[-1].cells
    total_cells[0].text = labels["total"] + f" ({labels['unit']})"
    total_cells[1].text = _fmt(boe.total.optimistic, locale)
    total_cells[2].text = _fmt(boe.total.likely, locale)
    total_cells[3].text = _fmt(boe.total.pessimistic, locale)

    doc.add_heading(labels["assumptions"], level=1)
    for assumption in boe.global_assumptions:
        doc.add_paragraph(assumption.text, style="List Bullet")
    for line in boe.lines:
        for assumption in line.assumptions:
            doc.add_paragraph(f"[{line.work_item_id}] {assumption.text}", style="List Bullet")

    doc.add_heading(labels["risks"], level=1)
    for risk in boe.global_risks:
        doc.add_paragraph(risk.text, style="List Bullet")
    for line in boe.lines:
        for risk in line.risks:
            suffix = (
                f" ({labels['contingency']}: {_fmt(risk.contingency_pd, locale)} {labels['unit']})"
                if risk.contingency_pd is not None
                else ""
            )
            doc.add_paragraph(f"[{line.work_item_id}] {risk.text}{suffix}", style="List Bullet")

    open_questions = [q for q in boe.questions if q.status not in ("answered", "applied")]
    if open_questions:
        doc.add_heading(labels["questions"], level=1)
        for question in open_questions:
            doc.add_paragraph(f"{question.id}: {question.question}", style="List Bullet")

    doc.add_heading(labels["provenance"], level=1)
    for line in boe.lines:
        doc.add_paragraph(f"{line.work_item_id} — {line.basis_note or ''}")
        for ref in line.evidence:
            doc.add_paragraph(f"  {ref.uri}" + (f" ({ref.label})" if ref.label else ""))

    doc.add_heading(labels["signatures"], level=1)
    sig_table = doc.add_table(rows=1 + max(len(boe.signatures), 2), cols=4)
    sig_table.style = "Table Grid"
    for column, text in enumerate(
        (labels["name"], labels["role"], labels["scope"], labels["date"])
    ):
        sig_table.rows[0].cells[column].text = text
    for row, signature in enumerate(boe.signatures, start=1):
        cells = sig_table.rows[row].cells
        cells[0].text = signature.name
        cells[1].text = signature.role
        cells[2].text = signature.scope
        # The signing date belongs in the artefact: an unsigned-looking column would
        # make the exported document unusable as the record it claims to be.
        cells[3].text = (
            signature.signed_at.strftime("%Y-%m-%d %H:%M") if signature.signed_at else ""
        )

    doc.save(str(path))
