"""Generate the synthetic Turkish BRD fixtures (.docx) and their planted-feature manifest.

Source of truth for fixtures/brd/ — edit the CONTENT definitions here and regenerate;
never hand-edit the .docx files (fixtures/README.md). Deterministic: no randomness, no
timestamps beyond the fictional dates embedded in content.

Usage:  python generate_brds.py            (writes into ../brd/)
Deps:   python-docx
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT_DIR = Path(__file__).resolve().parent.parent / "brd"


# ---------------------------------------------------------------- helpers


def _style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _meta_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, value) in enumerate(rows):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value


def _grid(doc: Document, header: list[str], body: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    for j, cell_text in enumerate(header):
        table.rows[0].cells[j].text = cell_text
    for i, row in enumerate(body, start=1):
        for j, cell_text in enumerate(row):
            table.rows[i].cells[j].text = cell_text


def _signature_block(doc: Document, names: list[tuple[str, str]]) -> None:
    doc.add_heading("Onay", level=1)
    _grid(doc, ["Ad Soyad", "Rol", "İmza / Tarih"], [[n, r, ""] for n, r in names])


# ---------------------------------------------------------------- BRD 01 — clean, templated


def brd_01(doc: Document) -> None:
    doc.add_heading("İş Gereksinim Dokümanı — Kampanya Bazlı Cihaz Taksitlendirme", level=0)
    _meta_table(
        doc,
        [
            ("Müşteri", "Aurora Telekom A.Ş."),
            ("Doküman No", "AUR-BRD-2026-041"),
            ("Tarih", "15.06.2026"),
            ("Versiyon", "1.2"),
            ("Hazırlayan", "S. Kaya — İş Analisti"),
        ],
    )

    doc.add_heading("1. Amaç", level=1)
    doc.add_paragraph(
        "Aurora Telekom, bireysel faturalı abonelerine kampanya bazlı cihaz "
        "taksitlendirme imkânı sunmayı hedeflemektedir. Bu doküman, taksitlendirme "
        "yeteneğinin iş gereksinimlerini üst seviyede tanımlar."
    )

    doc.add_heading("2. Kapsam", level=1)
    doc.add_paragraph(
        "Kapsam; kampanya tanımlama, taksit planı oluşturma, faturaya yansıtma ve bayi "
        "kanalından satış akışlarını içerir. Kurumsal segment ve ön ödemeli hat satışları "
        "bu fazın kapsamı dışındadır."
    )

    doc.add_heading("3. İş Gereksinimleri", level=1)
    for code, text in [
        (
            "G-01",
            (
                "Pazarlama ekibi, kampanya yönetimi ekranından taksitli cihaz "
                "kampanyası tanımlayabilmelidir (cihaz listesi, taksit sayısı seçenekleri, "
                "kampanya dönemi)."
            ),
        ),
        ("G-02", "Taksit planları 3, 6, 12 ve 24 ay seçenekleriyle sunulmalıdır."),
        ("G-03", "Taksit tutarı, abonenin aylık faturasına ayrı kalem olarak yansıtılmalıdır."),
        ("G-04", "Kurumsal müşteriler için farklı koşullar uygulanabilir."),
        ("G-05", "Bayi, satış anında abonenin taksitlendirmeye uygunluğunu sorgulayabilmelidir."),
        (
            "G-06",
            (
                "Kampanya bazında komisyon oranları aşağıdaki matrise göre "
                "uygulanmalıdır (bkz. Tablo 3.2-a)."
            ),
        ),
        ("G-07", "Cihaz iadesi durumunda taksit planı kapatılmalıdır."),
        ("G-08", "Taksitlendirme hareketleri günlük olarak muhasebe sistemine raporlanmalıdır."),
    ]:
        doc.add_paragraph(f"{code}: {text}", style="List Bullet")

    doc.add_heading("3.2 Taksit Komisyon Matrisi (Tablo 3.2-a)", level=2)
    _grid(
        doc,
        ["Taksit Sayısı", "Komisyon Oranı", "Asgari Sepet Tutarı"],
        [
            ["3 ay", "%2,4", "1.500 TL"],
            ["6 ay", "%4,1", "1.500 TL"],
            ["12 ay", "%7,8", "3.000 TL"],
            ["24 ay", "%12,5", "6.000 TL"],
        ],
    )

    doc.add_heading("4. Varsayımlar ve Kısıtlar", level=1)
    doc.add_paragraph("V-01: Bu iş için ayrılan bütçe azami 90 adam-gündür.", style="List Bullet")
    doc.add_paragraph(
        "V-02: Kampanya lansmanı 30 Eylül 2026 tarihine yetişmelidir.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "V-03: Cihaz stok yönetimi mevcut bayi portalı üzerinden yürütülecektir.",
        style="List Bullet",
    )

    _signature_block(
        doc,
        [
            ("S. Kaya", "İş Analisti (Aurora Telekom)"),
            ("A. Demir", "Analist (Meridyen Teknoloji)"),
        ],
    )


# ---------------------------------------------------------------- BRD 02 — clean, low ambiguity (control)


def brd_02(doc: Document) -> None:
    doc.add_heading("İş Gereksinim Dokümanı — Konsolide Fatura", level=0)
    _meta_table(
        doc,
        [
            ("Müşteri", "Aurora Telekom A.Ş."),
            ("Doküman No", "AUR-BRD-2026-052"),
            ("Tarih", "02.07.2026"),
            ("Versiyon", "1.0"),
            ("Hazırlayan", "E. Şahin — Ürün Yöneticisi"),
        ],
    )

    doc.add_heading("1. Amaç", level=1)
    doc.add_paragraph(
        "Aynı müşteri numarası altındaki birden çok aboneliğin tek bir konsolide fatura "
        "altında birleştirilmesi hedeflenmektedir."
    )

    doc.add_heading("2. İş Gereksinimleri ve Kabul Kriterleri", level=1)
    _grid(
        doc,
        ["No", "Gereksinim", "Kabul Kriteri"],
        [
            [
                "G-01",
                (
                    "Müşteri, çağrı merkezi veya online kanal üzerinden konsolide fatura "
                    "talebinde bulunabilmelidir."
                ),
                "Talep sonrası ilk fatura döneminde tek fatura üretilir.",
            ],
            [
                "G-02",
                "Konsolide faturada abonelik bazında alt kırılım gösterilmelidir.",
                "Her abonelik; sabit ücret, kullanım ve vergiler kırılımıyla listelenir.",
            ],
            [
                "G-03",
                "Konsolidasyon yalnızca aynı fatura dönemine sahip abonelikler için yapılmalıdır.",
                "Farklı dönemli abonelik eklenmek istendiğinde sistem yönlendirme mesajı gösterir.",
            ],
            [
                "G-04",
                "Konsolide faturadan çıkış (ayrıştırma) desteklenmelidir.",
                "Ayrıştırma bir sonraki fatura döneminde geçerli olur.",
            ],
            [
                "G-05",
                "Ödeme, konsolide fatura toplamı üzerinden tek seferde alınmalıdır.",
                "Kısmi ödeme mevcut kurallarla abonelik bazına dağıtılır.",
            ],
            [
                "G-06",
                "Konsolide fatura PDF'i mevcut şablon standardına uygun üretilmelidir.",
                "PDF, kurumsal şablon v4 ile birebir uyumludur.",
            ],
        ],
    )

    _signature_block(
        doc,
        [
            ("E. Şahin", "Ürün Yöneticisi (Aurora Telekom)"),
            ("A. Demir", "Analist (Meridyen Teknoloji)"),
        ],
    )


# ---------------------------------------------------------------- BRD 03 — messy, unstructured, high ambiguity


def brd_03(doc: Document) -> None:
    doc.add_heading("Bayi Sipariş Entegrasyonu hk.", level=0)
    doc.add_paragraph("Aurora Telekom A.Ş. — Saha Satış Direktörlüğü — Haziran 2026")

    doc.add_paragraph(
        "Merhaba, aşağıda özetlediğimiz konuyu bir sonraki sürüm planına almanızı rica "
        "ediyoruz. Bayilerimizin kullandığı sipariş sisteminin bizim tarafla entegre "
        "olması gerekiyor. Şu an bayiler siparişleri ayrı bir ekrandan tekrar giriyor, "
        "bu da hem zaman kaybı hem hata demek."
    )
    doc.add_paragraph(
        "Beklentimiz siparişin bayi tarafında oluştuğu anda bizim tarafta da otomatik "
        "açılması. Stok durumu da anlık görünsün istiyoruz. Ayrıca kampanyalı satışlarda "
        "onay adımı gerekiyorsa bu da akışın içinde olmalı. Sipariş iptallerinde iki "
        "taraf da güncellenmeli."
    )
    doc.add_paragraph(
        "Benzer bir entegrasyonu geçen yıl başka bir tedarikçiyle 40 günde "
        "tamamlamıştık; bu işin de benzer sürede biteceğini düşünüyoruz."
    )
    doc.add_paragraph(
        "Raporlama tarafında da günlük sipariş özetinin yönetim panosuna düşmesi lazım. "
        "Detayları toplantıda konuşuruz. Teşekkürler."
    )
    doc.add_paragraph("— K. Arslan, Saha Satış Direktörü")


# ---------------------------------------------------------------- BRD 04 — micro-CR (small-item floor scenario)


def brd_04(doc: Document) -> None:
    doc.add_heading("Değişiklik Talebi — Ön Ödemeliden Faturalıya Geçişte Bakiye Taşıma", level=0)
    _meta_table(
        doc,
        [
            ("Müşteri", "Aurora Telekom A.Ş."),
            ("Talep No", "AUR-CR-2026-118"),
            ("Tarih", "21.07.2026"),
            ("Talep Sahibi", "B. Çelik — Müşteri Deneyimi"),
        ],
    )

    doc.add_heading("Talep", level=1)
    doc.add_paragraph(
        "Ön ödemeli hattan faturalı tarifeye geçiş yapan abonenin kalan TL bakiyesi, "
        "geçiş sonrası ilk faturasına indirim kalemi olarak yansıtılmalıdır."
    )
    doc.add_heading("Not", level=1)
    doc.add_paragraph(
        "Bakiyenin ilk fatura tutarını aşması durumunda izlenecek yöntem işletme ile "
        "netleştirilecektir."
    )


# ---------------------------------------------------------------- manifest

MANIFEST = {
    "universe": "aurora",
    "generator": "tools/generate_brds.py",
    "brds": [
        {
            "file": "BRD-AUR-26-01-taksitlendirme.docx",
            "title": "Kampanya Bazlı Cihaz Taksitlendirme",
            "maturity": "clean-templated",
            "requirement_count": 8,
            "tables": 3,
            "planted_anchors": [
                {"type": "budget", "snippet": "azami 90 adam-gün"},
                {"type": "deadline", "snippet": "30 Eylül 2026"},
            ],
            "planted_ambiguities": [
                {
                    "ref": "G-04",
                    "issue": "undefined-segment-terms",
                    "note": "'Kurumsal müşteriler için farklı koşullar' — segment ve koşullar tanımsız",
                },
                {
                    "ref": "G-07",
                    "issue": "missing-acceptance-criteria",
                    "note": "İade senaryosunda tahsil edilmiş taksitlerin akıbeti belirsiz",
                },
            ],
            "expected_modules": [
                "campaign-engine",
                "billing-core",
                "dealer-portal",
                "invoice-render",
                "integration-hub",
            ],
        },
        {
            "file": "BRD-AUR-26-02-konsolide-fatura.docx",
            "title": "Konsolide Fatura",
            "maturity": "clean-templated",
            "requirement_count": 6,
            "tables": 3,
            "planted_anchors": [],
            "planted_ambiguities": [],
            "expected_modules": ["billing-core", "invoice-render", "crm-suite", "selfcare-web"],
        },
        {
            "file": "BRD-AUR-26-03-bayi-siparis-entegrasyonu.docx",
            "title": "Bayi Sipariş Entegrasyonu",
            "maturity": "messy-unstructured",
            "requirement_count": None,
            "tables": 0,
            "planted_anchors": [
                {"type": "analogy", "snippet": "geçen yıl başka bir tedarikçiyle 40 günde"},
            ],
            "planted_ambiguities": [
                {
                    "ref": None,
                    "issue": "no-structure",
                    "note": "Numarasız, e-posta üslubu; gereksinimler paragraf içinde",
                },
                {
                    "ref": None,
                    "issue": "missing-volumetrics",
                    "note": "Sipariş hacmi, bayi sistemi/protokol, SLA tanımsız",
                },
                {
                    "ref": None,
                    "issue": "undefined-approval-flow",
                    "note": "'Onay adımı gerekiyorsa' — koşul ve onaylayan tanımsız",
                },
            ],
            "expected_modules": ["dealer-portal", "integration-hub", "product-catalog"],
        },
        {
            "file": "BRD-AUR-26-04-bakiye-tasima.docx",
            "title": "Ön Ödemeliden Faturalıya Geçişte Bakiye Taşıma",
            "maturity": "micro-cr",
            "requirement_count": 1,
            "tables": 1,
            "planted_anchors": [],
            "planted_ambiguities": [
                {
                    "ref": None,
                    "issue": "unresolved-edge-case",
                    "note": "Bakiye > ilk fatura durumu bilinçli olarak açık bırakılmış",
                },
            ],
            "expected_modules": ["billing-core", "crm-suite"],
        },
    ],
}

BUILDERS = {
    "BRD-AUR-26-01-taksitlendirme.docx": brd_01,
    "BRD-AUR-26-02-konsolide-fatura.docx": brd_02,
    "BRD-AUR-26-03-bayi-siparis-entegrasyonu.docx": brd_03,
    "BRD-AUR-26-04-bakiye-tasima.docx": brd_04,
}


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for filename, build in BUILDERS.items():
        doc = Document()
        _style(doc)
        build(doc)
        doc.save(OUT_DIR / filename)
        print(f"wrote {filename}")
    manifest_path = OUT_DIR / "manifest.json"
    manifest_path.write_text(
        json.dumps(MANIFEST, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    print(f"wrote {manifest_path.name}")


if __name__ == "__main__":
    main()
